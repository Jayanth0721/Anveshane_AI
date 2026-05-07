import os
from docx import Document

os.makedirs("docs for tender2", exist_ok=True)

# 1. Tender 2 Document
tender_doc = Document()
tender_doc.add_heading('Tender Document: Tender 2', 0)

tender_doc.add_paragraph("Tender Title: Tender 2")
tender_doc.add_paragraph("Description: Govt Scool Uniform")
tender_doc.add_paragraph("Sector: Education")
tender_doc.add_paragraph("Investment Amount: 10cr")
tender_doc.add_paragraph("Duration (days): 60")
tender_doc.add_paragraph("Department: Children Welfare")
tender_doc.add_paragraph("Work Location: Karnataka")
tender_doc.add_paragraph("Penalty per day (delay): 20000")
tender_doc.add_paragraph("Application closing date: 10-05-2026 06:01")
tender_doc.add_paragraph("Max penalty period (days, default 180): 180")

tender_doc.save("docs for tender2/Tender 2.docx")

# 2. Eligible Bidder Document
eligible_doc = Document()
eligible_doc.add_heading('Bidder Submission', 0)

eligible_doc.add_paragraph("Company Name: M/s. Education Suppliers Ltd.")
eligible_doc.add_paragraph("Bidder Name: John Doe")
eligible_doc.add_paragraph("Subject: Proposal for Govt School Uniform Tender")

eligible_doc.add_heading('Company Profile & Eligibility', level=1)
eligible_doc.add_paragraph("We are a leading supplier in the Education sector, specializing in student uniforms, books, and sports kits.")
eligible_doc.add_paragraph("Turnover: INR 50 Crore for the last financial year.")
eligible_doc.add_paragraph("Certifications: GSTIN 29AAAAA0000A1Z5, ISO 9001:2015 certified company.")

eligible_doc.add_heading('Project References', level=1)
eligible_doc.add_paragraph("References: We have successfully completed work orders for 100,000 Govt School Uniforms for the Karnataka Education Department.")

eligible_doc.add_heading('Commercial Proposal', level=1)
eligible_doc.add_paragraph("Based on the requirements for the Govt School Uniforms tender, we submit our commercial proposal.")
eligible_doc.add_paragraph("Quoted Amount: INR 8 Crore")
eligible_doc.add_paragraph("Completion timeline: 45 days to deliver the complete order.")

eligible_doc.save("docs for tender2/bidder_eligible.docx")

# 3. Not Eligible Bidder Document
not_eligible_doc = Document()
not_eligible_doc.add_heading('Bidder Submission', 0)

not_eligible_doc.add_paragraph("Company Name: M/s. Fast Builders.")
not_eligible_doc.add_paragraph("Bidder Name: Jane Smith")
not_eligible_doc.add_paragraph("Subject: Proposal for Govt School Uniform Tender")

not_eligible_doc.add_heading('Company Profile & Eligibility', level=1)
not_eligible_doc.add_paragraph("We are a leading construction company in the Infrastructure sector. We do road construction, drainage, and bridge building.")
# Intentionally omitting turnover and certifications

not_eligible_doc.add_heading('Commercial Proposal', level=1)
not_eligible_doc.add_paragraph("We can attempt to source uniforms as well.")
not_eligible_doc.add_paragraph("Quoted Amount: INR 15 Crore")
not_eligible_doc.add_paragraph("Completion timeline: 120 days")

not_eligible_doc.save("docs for tender2/bidder_not_eligible.docx")

print("Successfully created docs in 'docs for tender2'")
